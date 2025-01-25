import streamlit as st
from utils.config import *
from parser import extractor, extract_trademark_details, extract_search_target
from comparator import *
from models import replace_disallowed_words
import requests
import json
from io import BytesIO
import logging
import fitz
import uuid
import asyncio
import json
from docx import Document


if "documents" not in st.session_state:
    st.session_state.documents = {}
if "removed_documents" not in st.session_state:
    st.session_state.removed_documents = []
if "session_id" not in st.session_state:
    st.session_state.session_id = str(uuid.uuid4())

st.title("AttorneyAI")

uploaded_files = st.file_uploader(
    "Upload files less than 400 pages",
    type=["pdf", "docx", "xlsx", "pptx"],
    accept_multiple_files=True,
    help="If your question is not answered properly or there's an error, consider uploading smaller documents or splitting larger ones.",
    label_visibility="collapsed",
)

combined_responses = ""
comparison_results = {
    "High": [],
    "Moderate": [],
    "Name-Match": [],
    "Low": [],
    "No-conflict": [],
}

if uploaded_files:
    new_files = []
    for uploaded_file in uploaded_files:
        if (
            uploaded_file.name
            not in [
                st.session_state.documents[doc_id]["name"]
                for doc_id in st.session_state.documents
            ]
            and uploaded_file.name not in st.session_state.removed_documents
        ):
            new_files.append(uploaded_file)
    extracted_details = []
    docu = Document()
    for new_file in new_files:
        st.success(f"File Selected: {new_file.name}")
        pdf_bytes = new_file.read()
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        extracted_pages = []
        for page_num, page in enumerate(doc, start=1):
            # Extract text with optional clipping
            text = page.get_text()
            text = replace_disallowed_words(text)
            extracted_pages.append(text)
        target_search = extract_search_target(doc)
        if target_search:
            docu.add_heading("Target search:")
            docu.add_paragraph(f"Proposed trademark: {target_search["mark_searched"]}\nSearch classes: {target_search["target_class"]}\nSearch Goods/Services: {target_search["goods_services"]}")
            docu.add_heading("Details extracted")
        st.write(target_search)
        Index, document = extractor(doc)
        if Index.startswith("```json"):
            Index = Index[7:]  # Remove the "```json" part
        if Index.endswith("```"):
            Index = Index[:-3]  # Remove the trailing "```"
        
        # Convert the remaining text to a Python object
        try:
            Index = json.loads(Index)
        except:
            print("Error in json parsing:")

        st.write(Index)
            
        async def parallel_extraction():
            semaphore = asyncio.Semaphore(10)  # Limit to 10 concurrent tasks
            tasks = []
            for i in range(len(Index)):
                start_page = int(Index[i]["page-start"]) - 1
                if i == len(Index) - 1:
                    end_page = start_page + 4
                else:
                    end_page = int(Index[i + 1]["page-start"]) - 1
        
                document_chunk = "\n".join(extracted_pages[start_page:end_page])
                tasks.append(
                    extract_trademark_details(document_chunk, Index[i]["name"], target_search['mark_searched'], semaphore)
                )
        
            return await asyncio.gather(*tasks)

        comparison_results = {
            "High": [],
            "Moderate": [],
            "Name-Match": [],
            "Low": [],
            "No-conflict": [],
        }
        async def process_trademarks():
            extracted_details = await parallel_extraction()
            count = len(extracted_details)
            docu.add_paragraph(f"Total no. of trademarks extracted: {count}")
            proposed_name = target_search["mark_searched"]
            proposed_class = target_search["target_class"]
            proposed_goods_services = target_search["goods_services"]

            for details in extracted_details:
                if not details or "error" in details:
                    continue
                st.write(details)
                docu.add_paragraph(f"Trademark name: {details["trademark_name"]}\nStatus: {details["status"]}\nSerial/Registration Number: {details["serial_number"]}/{details["registration_number"]}\nInternational Class: {details["international_class_number"]}\nGoods/Services: {details["goods_services"]}\nOwner: {details["owner"]}\nFiled Date: {details["filed_date"]}\nDesign Phrase: {details["design_phrase"]}")
                comparision_result = compare_trademarks(
                    details, proposed_name, proposed_class, proposed_goods_services
                )

                conflict_grade = comparision_result.get("conflict_grade")
                comparison_results[conflict_grade].append(comparision_result)

        asyncio.run(process_trademarks())

    
    for conflict_grade, results in comparison_results.items():
        count = len(results)
        docu.add_paragraph(f"{conflict_grade}: {count} entries")

    for conflict_grade, results in comparison_results.items():
        if results:
            docu.add_heading(conflict_grade, level=2)
            table = docu.add_table(rows=1, cols=5)
            header = table.rows[0].cells
            header[0].text = "Trademark Name and Class Number"
            header[1].text = "Trademark Status"
            header[2].text = "Serial/Registration Number"
            header[3].text = "Owner Name"
            header[4].text = "Design/Word"

            for result in results:
                row = table.add_row().cells
                row[0].text = (
                    f"{result['Trademark name']} (Class {result['Trademark class Number']})"
                )
                row[1].text = result["Trademark status"]
                row[2].text = (
                    f"{result['Trademark serial number']} / {result['Trademark registration number']}"
                )
                row[3].text = result["Trademark owner"]
                row[4].text = (
                    "Design" if result["Trademark design phrase"] else "Word"
                )
    for conflict_grade, results in comparison_results.items():
        for result in results:
            docu.add_paragraph(
                f"{result['Trademark name']} reasoning: {result['reasoning']}\n"
            )


    #for details in extracted_details:
        

    output = BytesIO()
    docu.save(output)
    output.seek(0)

    # Streamlit app logic
    st.title("Trademark Conflict Analysis")
    st.write("Download the trademark conflict analysis document below.")

    # Add a download button
    st.download_button(
        label="Download Analysis Document",
        data=output,
        file_name="Trademark_Conflict_Analysis.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
